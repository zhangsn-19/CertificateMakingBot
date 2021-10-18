from docx.shared import Inches, Pt
import csv
import requests
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement

def addShadeToTxt(run,color="FFAAAA"):

# Get the XML tag
    tag = run._r
    #print(run.element.xml)

# Create XML element
    shd = OxmlElement('w:shd')

# Add attributes to the element
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)

# This is the tricky part
    run.font.size = Pt(14)
    #define debug
    #print(run.element.xml)

    tag.rPr.append(shd)
from docx.enum.text import WD_ALIGN_PARAGRAPH

# docx 处理
from docx import Document
Doc = []
distance = Inches(3)

def addbreak(obj,j):
    i=0
    while i<=j:
        obj.add_paragraph(" ")
        obj.paragraphs[0].runs[0].add_break
        i=i+1

fontDefault = u'宋体'
def chg_font(obj,fontname=fontDefault,size=None):
    obj.font.name = fontname
    obj._element.rPr.rFonts.set(qn('w:eastAsia'),fontname)
 
    if size and isinstance(size,Pt):
        obj.font.size = size
def chg_doc(obj,name):
    distance = Inches(2)
    sec = obj.sections[0]             # sections对应文档中的“节”
    sec.left_margin = distance     # 以下依次设置左、右、上、下页面边距
    sec.right_margin = distance
    sec.top_margin = distance
    sec.bottom_margin = distance
    sec.page_width = Inches(15)        #设置页面宽度
    sec.page_height = Inches(20)       #设置页面高度
    ##设置默认字体
    chg_font(obj.styles['Normal'],fontname='宋体',size=36)
 
 
   # i=0
   ## while i<=10:
     ##   obj.add_paragraph(" ")
       ## obj.paragraphs[0].runs[0].add_break
        ##i=i+1
    
    addbreak(obj,5)

    obj.add_heading('小伙伴计划早安打卡回忆',0)
    addbreak(obj,5)
    
    p1=obj.add_paragraph()
    run1 = p1.add_run('““一日之计在于晨”，每天早晨在清新的空气与温暖的阳光中苏醒，带着新一天的元气与活力完成一件小事。不知不觉，你已经伴着每日一事的节奏，发现平凡生活中的微小光亮，将生活编织成了一首沁人心脾的诗。在与打卡相伴的日子里，每一天的点滴收获都促使你变成更优秀的自己。接下来，让我们翻开这本纪念册，品味这独特又珍贵的回忆吧。 ”') 

    run1.font.name = '微软雅黑'
    run1.element.rPr.rFonts.set(qn('w:eastAsia'),u'微软雅黑')
    run1.font.size = Pt(21)
    

    section = obj.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #此处可以修改 卷数和对应名称
    paragraph.text = "小伙伴计划回忆集 第1卷 早安世界" + name
 


import urllib.request

cnt = 0
totalCnt = 0
def add_text(content):
    flag = 0
    for i in Doc:
        if i["Name"] is content["name"]:
            DDoc = i["Docx"]
            flag = 1
    if flag is 0:
        doc = Document()
        docs = {"Docx":doc,"Name":content["name"]}
        Doc.append(docs)
    DDoc = docs["Docx"]
    chg_doc(DDoc,content["name"])
    section = doc.add_section()
    section._sectPr.xpath('./w:cols')[0].set(qn('w:num'),'2')
    #print(content)
    global cnt
    global totalCnt
    print(totalCnt)
    totalCnt += 1
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  
    run = paragraph.add_run(content['name'])
    addShadeToTxt(run)
    run.bold = True #设置字体为粗体
    
    chg_font(run,fontname='微软雅黑', size=Pt(14))  #设置字体和字

    for concreteContent in content['activity']:
        addShadeToTxt(run)
        run = paragraph.add_run(concreteContent[0])
        chg_font(run,fontname='微软雅黑', size=Pt(11))
        run = paragraph.add_run(concreteContent[1])
        run.add_break()
        chg_font(run,fontname='微软雅黑', size=Pt(12))
        images = concreteContent[2].split(',')
        for image in images:
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  
            run = paragraph.add_run("")
            cnt += 1
            if cnt % 2 == 0:
                imagename = str(cnt) + ".jpg"
                image = image.strip()
                if ".jpg" in image:
                    #print(image)
                    urllib.request.urlretrieve(image,filename=imagename)
                    run.add_picture(imagename,width=Inches(4.25))
            
    ph_format =paragraph.paragraph_format
    
    ph_format.space_before =Pt(10)     #设置段前间距
    
    ph_format.space_after =Pt(12)       #设置段后间距
    
    ph_format.line_spacing=Pt(19)       #设置行间距

# the '1.csv' can be replaced with your file name as long as it is in the same directory with this python file

f = csv.reader(open('1.csv', 'r'))
l = []
Flag = 0
for i in f:
    if not Flag:
        Flag = 1
    else:
        name = i[0]
        flag = 0
        for u in l:
            if u['name'] == name:
                flag = 1
                u['activity'].append((i[5], i[7], i[10]))
                break
        if flag == 0:
            l.append({"name": name, "activity": [(i[5], i[7], i[10]), ]})

#调试
#for r in l:
#   print(r)

lCnt = 0
for thing in l:
    if lCnt < 20:
        add_text(thing)
        for docs in Doc:
            #define debug
            # print(docs["Name"])
            #
            docs["Docx"].save(docs["Name"]+'.docx')
        lCnt += 1
import os 
for i in range(1,cnt//2):
    r = str(i*2) + ".jpg"
    if os.path.exists(r):
        os.remove(r)
